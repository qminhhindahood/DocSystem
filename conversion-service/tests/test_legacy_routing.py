"""test_legacy_routing.py — LEGACY_TEXT routing through the stack (ticket 04).

Contract (ticket 04 acceptance):
1. triage_page classifies decodable legacy-encoded text as LEGACY_TEXT —
   not SCANNED (no Gemini key needed; table decode is lossless where OCR
   never can be) and not DIGITAL_TEXT (raw mojibake must not flow into the
   rule engine).
2. Undecodable garbage stays SCANNED (vision fallback unchanged).
3. pipeline converts a LEGACY_TEXT page losslessly: DOCX contains the
   decoded Unicode phrase; page_types records "LEGACY_TEXT"; extracted_chars
   counted from the decoded text (fidelity ledger consistency).
4. main.py's scanned-admission gate: legacy-decodable PDF passes without
   a Gemini key; genuinely blank/scanned PDF still trips it.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import fitz
from docx import Document

from triage.triage import DIGITAL_TEXT, LEGACY_TEXT, SCANNED, triage_page
from tests.test_triage import FakePage

sys.path.insert(0, str(Path(__file__).resolve().parent))
FIXTURES = json.loads(
    (Path(__file__).parent / "fixtures_legacy.json").read_text(encoding="utf-8")
)


def _make_legacy_pdf(path: str, raw: str) -> str:
    """Insert raw legacy bytes as the text layer of a fresh PDF."""
    doc = fitz.open()
    page = doc.new_page()
    font = fitz.Font(fontname="helv")
    tw = fitz.TextWriter(page.rect)
    tw.append((72, 96), raw, font=font, fontsize=14)
    tw.write_text(page)
    doc.save(path)
    doc.close()
    return path


def _docx_text(path: str) -> str:
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestTriageLegacyRouting:
    def test_tcvn3_garbage_routes_to_legacy_text(self):
        assert triage_page(FakePage(text=FIXTURES["tcvn3_phrase"])) == LEGACY_TEXT

    def test_vni_garbage_routes_to_legacy_text(self):
        assert triage_page(FakePage(text=FIXTURES["vni_phrase"])) == LEGACY_TEXT

    def test_healthy_vietnamese_stays_digital(self):
        # Health-gain guard: healthy Unicode pages never take the legacy
        # path even though TCVN3 keys overlap Latin-1 codepoints.
        # (>50 letters: triage requires MIN_TRUSTED_LETTERS before trusting
        # any text layer — short slivers route to vision.)
        vn = (
            "Cộng hòa xã hội chủ nghĩa Việt Nam. Độc lập - Tự do - Hạnh phúc. "
            "Căn cứ Nghị định số 30/2020/NĐ-CP ngày 05 tháng 3 năm 2020 của "
            "Chính phủ về công tác văn thư; xét đề nghị của Cục trưởng Cục "
            "Văn thư và Lưu trữ nhà nước về việc ban hành quy định này."
        )
        assert triage_page(FakePage(text=vn)) == DIGITAL_TEXT

    def test_undecodable_garbage_stays_scanned(self):
        garbage = "CNG HA X HI CH NGHA VIT NAM c lp T do Hnh phc " * 5
        assert triage_page(FakePage(text=garbage)) == SCANNED

    def test_legacy_page_with_fullpage_image_stays_scanned(self):
        # A dominant scan image is the stronger signal: vision remains the
        # only honest path; a decodable text sliver must not reclassify it.
        page = FakePage(
            text=FIXTURES["tcvn3_phrase"],
            images=[(42, 0, 0, 0, 0, 0, 0)],
            image_rects={42: [type("R", (), {"width": 590, "height": 840})()]},
        )
        assert triage_page(page) == SCANNED


class TestPipelineLegacyBranch:
    def test_convert_pdf_legacy_page_lossless(self, tmp_path):
        from pipeline import convert_pdf

        pdf = _make_legacy_pdf(str(tmp_path / "legacy.pdf"), FIXTURES["tcvn3_phrase"])
        out = str(tmp_path / "legacy.docx")
        path, report = convert_pdf(pdf, out)

        assert report.page_types.get("LEGACY_TEXT", 0) >= 1
        text = _docx_text(path)
        assert FIXTURES["tcvn3_expected"] in text
        # fidelity ledger: extracted chars counted from decoded text
        assert report.extracted_chars >= len(FIXTURES["tcvn3_expected"])

    def test_convert_pdf_vni_page_lossless(self, tmp_path):
        from pipeline import convert_pdf

        pdf = _make_legacy_pdf(str(tmp_path / "vni.pdf"), FIXTURES["vni_phrase"])
        out = str(tmp_path / "vni.docx")
        path, report = convert_pdf(pdf, out)

        assert report.page_types.get("LEGACY_TEXT", 0) >= 1
        assert FIXTURES["vni_expected"] in _docx_text(path)

    def test_convert_pdf_undecodable_degrades_honestly(self, tmp_path):
        from pipeline import convert_pdf

        pdf = _make_legacy_pdf(
            str(tmp_path / "garbage.pdf"),
            "CNG HA X HI CH NGHA VIT NAM c lp " * 5,
        )
        out = str(tmp_path / "garbage.docx")
        _, report = convert_pdf(pdf, out)

        assert report.page_types.get("LEGACY_TEXT", 0) == 0
        assert report.page_types.get("SCANNED", 0) >= 1


class TestMainGate:
    def test_legacy_pdf_admitted_without_gemini_key(self, tmp_path):
        from main import _has_scanned_pages

        pdf = _make_legacy_pdf(str(tmp_path / "legacy.pdf"), FIXTURES["tcvn3_phrase"])
        assert _has_scanned_pages(pdf) is False

    def test_blank_pdf_still_trips_gate(self, tmp_path):
        from main import _has_scanned_pages

        doc = fitz.open()
        doc.new_page()
        pdf = str(tmp_path / "blank.pdf")
        doc.save(pdf)
        doc.close()
        assert _has_scanned_pages(pdf) is True
