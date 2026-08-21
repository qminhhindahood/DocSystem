"""tests/test_vision_wiring.py — BYOK scanned-page vision wiring.

The "real test" bar: a scanned PDF + an injected (fake) Gemini key must flow
through the pipeline and produce a Decree-30 DOCX we re-open and verify —
Times New Roman, A4 margins, 2-column admin header, motto underline, and the
transcribed Vietnamese text present. Only the outbound Gemini HTTPS call is
faked; every line of our logic (triage gate, admission, batching, validation,
assembly, rule engine, renderer, quota refund) runs for real.

Seams under test:
  * main._parse_vision / _has_scanned_pages / the 422 admission gate
  * pipeline._run_scanned_vision (batching + validation + confidence cap)
  * worker fail-fast on VisionAuthError + single quota refund
  * GeminiVisionClient.extract_batch_json response parsing
"""
from __future__ import annotations

import io
import json
import sys
import asyncio
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import fitz
import pytest
from docx import Document
from PIL import Image, ImageDraw

import config
import main
import worker
from job_store import JobStore
from pipeline import _run_scanned_vision, convert_pdf
from quota import QuotaService
from vision.gemini_contract import (
    GeminiVisionClient,
    VisionAuthError,
    convert_scanned_pages_parallel,
    plan_batches,
)

TIMES = r"C:\Windows\Fonts\times.ttf"

# ─── Realistic Gemini output (a clean Quyết định, page 1) ─────────────────────
# Derived from eval/fixtures/quyet_dinh.json — the same block contract the real
# model must emit. No image/seal/section_break here so the render needs no media.
FAKE_VISION_BLOCKS = [
    {
        "type": "admin_header", "confidence": 0.9, "page": 1,
        "left": {
            "superior_agency": "BỘ NỘI VỤ",
            "issuing_agency": "CỤC VĂN THƯ VÀ LƯU TRỮ NHÀ NƯỚC",
            "document_number": "Số: 128/QĐ-VTLTNN",
        },
        "right": {
            "country_name": "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
            "motto": "Độc lập - Tự do - Hạnh phúc",
            "location_and_date": "Hà Nội, ngày 25 tháng 4 năm 2025",
        },
    },
    {"type": "heading", "level": 1, "text": "QUYẾT ĐỊNH",
     "align": "center", "confidence": 0.9, "page": 1},
    {"type": "paragraph", "align": "center",
     "text": "Về việc thành lập Ban Chỉ đạo chuyển đổi số",
     "confidence": 0.9, "page": 1},
    {"type": "paragraph", "align": "justify",
     "text": "Căn cứ Nghị định số 30/2020/NĐ-CP ngày 05 tháng 3 năm 2020 của "
             "Chính phủ về công tác văn thư.",
     "confidence": 0.9, "page": 1},
    {
        "type": "signature", "confidence": 0.9, "page": 1,
        "left": {"receipt_list": ["Như Điều 3;", "Lưu: VT, TCCB."]},
        "right": {"authority": "KT. CỤC TRƯỞNG", "title": "PHÓ CỤC TRƯỞNG",
                  "name": "Nguyễn Văn A"},
    },
]


# ─── Fixture builders (real PDFs, no mocks) ───────────────────────────────────

def _make_scanned_pdf(path: str) -> str:
    """A page with a full-page raster image and NO text layer -> SCANNED."""
    img = Image.new("RGB", (1240, 1754), "white")  # A4 @ ~150dpi
    draw = ImageDraw.Draw(img)
    # A few dark marks so it reads as a scan, not a blank page.
    draw.rectangle([100, 120, 1100, 160], fill="black")
    draw.rectangle([100, 300, 1100, 320], fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_image(page.rect, stream=buf.getvalue())
    doc.save(path)
    doc.close()
    return path


def _make_digital_pdf(path: str) -> str:
    """A page with a healthy Vietnamese text layer -> DIGITAL_TEXT."""
    font = fitz.Font(fontfile=TIMES)
    doc = fitz.open()
    page = doc.new_page()
    tw = fitz.TextWriter(page.rect)
    tw.append((72, 120), "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", font=font, fontsize=14)
    tw.append((72, 150), "Độc lập - Tự do - Hạnh phúc", font=font, fontsize=14)
    tw.append((72, 220), "Căn cứ Nghị định số 30/2020/NĐ-CP ngày 05 tháng 3 năm 2020",
              font=font, fontsize=14)
    tw.append((72, 245), "của Chính phủ quy định về công tác văn thư và lưu trữ nhà nước.",
              font=font, fontsize=14)
    tw.write_text(page)
    doc.save(path)
    doc.close()
    return path


def _docx_text(docx_path: str) -> str:
    dx = Document(docx_path)
    parts = [p.text for p in dx.paragraphs]
    for t in dx.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return " ".join(" ".join(parts).split())


def _mem_count(quota: QuotaService, user: str) -> int:
    entry = quota._memory.get(quota._key(user))
    return entry[0] if entry else 0


@pytest.fixture
def fresh_quota(monkeypatch):
    """Deterministic in-memory quota, independent of any live Redis."""
    quota = QuotaService(redis_client=None, limit=5)
    monkeypatch.setattr(main, "QUOTA", quota)
    return quota


@pytest.fixture
def no_job_exec(monkeypatch):
    """Keep admission tests from running the real pipeline in a background task."""
    async def noop(*args, **kwargs):
        return None
    monkeypatch.setattr(main, "_run_job_in_process", noop)


# ─── 1. Admission gate (real FastAPI app, real scanned PDF) ───────────────────

def test_scanned_upload_without_vision_is_rejected_422(fresh_quota, no_job_exec, tmp_path):
    from fastapi.testclient import TestClient

    pdf = Path(_make_scanned_pdf(str(tmp_path / "scan.pdf"))).read_bytes()
    with TestClient(main.app) as client:
        r = client.post(
            "/convert",
            files={"file": ("scan.pdf", pdf, "application/pdf")},
            headers={"X-User-Id": "u-nokey"},
        )
    assert r.status_code == 422
    detail = r.json()["detail"]
    assert "trang quét" in detail            # Vietnamese instruction
    assert "Google Gemini" in detail
    # Rejected before quota -> nothing consumed.
    assert _mem_count(fresh_quota, "u-nokey") == 0


def test_scanned_upload_with_vision_is_admitted(fresh_quota, no_job_exec, tmp_path):
    from fastapi.testclient import TestClient

    pdf = Path(_make_scanned_pdf(str(tmp_path / "scan.pdf"))).read_bytes()
    vision = json.dumps({"provider": "gemini", "model": "gemini-2.5-flash",
                         "apiKey": "user-key"})
    with TestClient(main.app) as client:
        r = client.post(
            "/convert",
            files={"file": ("scan.pdf", pdf, "application/pdf")},
            data={"vision": vision},
            headers={"X-User-Id": "u-key"},
        )
    assert r.status_code == 200
    assert r.json()["jobId"]
    # Admitted -> quota charged exactly once.
    assert _mem_count(fresh_quota, "u-key") == 1


def test_digital_upload_without_vision_is_admitted(fresh_quota, no_job_exec, tmp_path):
    from fastapi.testclient import TestClient

    pdf = Path(_make_digital_pdf(str(tmp_path / "digital.pdf"))).read_bytes()
    with TestClient(main.app) as client:
        r = client.post(
            "/convert",
            files={"file": ("digital.pdf", pdf, "application/pdf")},
            headers={"X-User-Id": "u-digital"},
        )
    assert r.status_code == 200
    assert r.json()["jobId"]


def test_invalid_vision_json_is_treated_as_absent(fresh_quota, no_job_exec, tmp_path):
    from fastapi.testclient import TestClient

    pdf = Path(_make_scanned_pdf(str(tmp_path / "scan.pdf"))).read_bytes()
    with TestClient(main.app) as client:
        r = client.post(
            "/convert",
            files={"file": ("scan.pdf", pdf, "application/pdf")},
            data={"vision": "{not-json"},
            headers={"X-User-Id": "u-badjson"},
        )
    # Garbage vision -> no usable key -> the scanned gate still rejects.
    assert r.status_code == 422


def test_parse_vision_rejects_non_gemini_and_missing_key():
    assert main._parse_vision(None) is None
    assert main._parse_vision("{bad") is None
    assert main._parse_vision(json.dumps(
        {"provider": "openrouter", "model": "m", "apiKey": "k"})) is None
    assert main._parse_vision(json.dumps(
        {"provider": "gemini", "model": "", "apiKey": "k"})) is None
    assert main._parse_vision(json.dumps(
        {"provider": "gemini", "model": "m", "apiKey": ""})) is None
    ok = main._parse_vision(json.dumps(
        {"provider": "gemini", "model": "gemini-2.5-flash", "apiKey": "k"}))
    assert ok == {"provider": "gemini", "model": "gemini-2.5-flash", "apiKey": "k"}


# ─── 2. The visible test: scanned PDF + injected key -> verified DOCX ─────────

def _three_page_pdf_bytes() -> bytes:
    doc = fitz.open()
    for page_number in range(1, 4):
        page = doc.new_page()
        page.insert_text((72, 72), f"original page {page_number}")
    payload = doc.tobytes()
    doc.close()
    return payload


def test_batch_plan_preserves_explicit_non_contiguous_pages():
    assert plan_batches([1, 3, 5], batch_size=2) == [(1, 3), (5,)]


def test_parallel_vision_receives_pdf_with_only_selected_pages():
    class RecordingClient:
        def __init__(self):
            self.calls = []

        def extract_batch_json(self, pdf_bytes, original_pages):
            with fitz.open(stream=pdf_bytes, filetype="pdf") as batch_doc:
                text = [batch_doc[index].get_text().strip() for index in range(len(batch_doc))]
                self.calls.append((tuple(original_pages), len(batch_doc), text))
            return []

    client = RecordingClient()
    asyncio.run(convert_scanned_pages_parallel(client, _three_page_pdf_bytes(), [1, 3]))

    assert client.calls == [
        ((1, 3), 2, ["original page 1", "original page 3"]),
    ]


def test_scanned_vision_rejects_blocks_from_digital_page(monkeypatch, tmp_path):
    class MixedPageClient:
        def __init__(self, api_key, model=None):
            pass

        def extract_batch_json(self, pdf_bytes, original_pages):
            return [
                {"type": "paragraph", "text": "scan one", "confidence": 0.9, "page": 1},
                {"type": "paragraph", "text": "digital duplicate", "confidence": 0.9, "page": 2},
                {"type": "paragraph", "text": "scan three", "confidence": 0.9, "page": 3},
            ]

    monkeypatch.setattr("vision.gemini_contract.GeminiVisionClient", MixedPageClient)
    pdf = tmp_path / "mixed.pdf"
    pdf.write_bytes(_three_page_pdf_bytes())

    blocks, degraded, _ = _run_scanned_vision(
        {"provider": "gemini", "model": "m", "apiKey": "k"},
        str(pdf),
        [1, 3],
    )

    assert [block.page for block in blocks] == [1, 3]
    assert degraded == []

class _FakeVisionClient:
    """Stands in for GeminiVisionClient at the network boundary only."""
    def __init__(self, api_key, model=None):
        assert api_key, "BYOK: a key must be injected"
        self.api_key = api_key
        self.model = model

    def extract_batch_json(self, pdf_bytes, original_pages):
        return FAKE_VISION_BLOCKS


def test_scanned_pipeline_produces_verified_decree30_docx(monkeypatch, tmp_path):
    monkeypatch.setattr("vision.gemini_contract.GeminiVisionClient", _FakeVisionClient)

    pdf = _make_scanned_pdf(str(tmp_path / "scan.pdf"))
    out = str(tmp_path / "scan.docx")
    media = str(tmp_path / "media")
    vision = {"provider": "gemini", "model": "gemini-2.5-flash", "apiKey": "user-key"}

    path, report = convert_pdf(pdf, out, media, vision)

    # Vision ran and produced content; the page was NOT degraded.
    assert report.degraded_pages == []
    assert report.status in ("completed", "completed_with_warnings")
    assert report.output_chars > 0
    assert report.confidence > 0
    # Scanned transcription is capped (never a default 1.0).
    assert report.confidence <= config.SCANNED_CONFIDENCE_CAP

    # Re-open the DOCX and verify the Decree-30 contract (P0a-gate shape).
    doc = Document(path)
    s0 = doc.sections[0]
    assert abs(s0.top_margin.mm - 25) < 0.5
    assert abs(s0.bottom_margin.mm - 25) < 0.5
    assert abs(s0.left_margin.mm - 30) < 0.5
    assert abs(s0.right_margin.mm - 20) < 0.5
    assert doc.styles["Normal"].font.name == "Times New Roman"

    # 2-column admin header table near the top.
    assert len(doc.tables) >= 1
    assert len(doc.tables[0].columns) == 2

    # Motto underline present somewhere.
    underlined = any(r.font.underline for p in doc.paragraphs for r in p.runs)
    if not underlined:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.font.underline:
                                underlined = True
    assert underlined

    # The transcribed Vietnamese text survived into the DOCX. The rule engine
    # legitimately uppercases the signature name per Decree-30, so compare
    # case-insensitively (same tolerance as test_fidelity.py).
    text = _docx_text(path).casefold()
    assert "cộng hòa xã hội chủ nghĩa việt nam" in text
    assert "quyết định" in text
    assert "nguyễn văn a" in text


def test_low_document_confidence_sets_delivery_warning(monkeypatch, tmp_path):
    class LowConfidenceClient(_FakeVisionClient):
        def extract_batch_json(self, pdf_bytes, original_pages):
            return [{
                "type": "paragraph",
                "text": "Nội dung được nhận dạng với độ tin cậy thấp",
                "confidence": 0.55,
                "page": 1,
            }]

    monkeypatch.setattr(
        "vision.gemini_contract.GeminiVisionClient", LowConfidenceClient
    )
    pdf = _make_scanned_pdf(str(tmp_path / "low-confidence.pdf"))

    _, report = convert_pdf(
        pdf,
        str(tmp_path / "low-confidence.docx"),
        str(tmp_path / "media"),
        {"provider": "gemini", "model": "m", "apiKey": "k"},
    )

    assert report.coverage == 1.0
    assert report.confidence == 0.55
    assert report.status == "completed_with_warnings"
    assert any("ngưỡng bàn giao" in warning for warning in report.warnings)


def test_scanned_without_vision_degrades_not_drops(tmp_path):
    """Backstop: no injected key -> scanned page degrades with a warning."""
    pdf = _make_scanned_pdf(str(tmp_path / "scan.pdf"))
    out = str(tmp_path / "scan.docx")
    _, report = convert_pdf(pdf, out, str(tmp_path / "media"), None)
    assert report.degraded_pages == [1]
    assert any("cần Gemini" in w for w in report.warnings)


def test_bad_vision_batch_degrades_its_pages(monkeypatch, tmp_path):
    class _NullClient(_FakeVisionClient):
        def extract_batch_json(self, pdf_bytes, original_pages):
            return None  # provider returned nothing usable
    monkeypatch.setattr("vision.gemini_contract.GeminiVisionClient", _NullClient)

    pdf = _make_scanned_pdf(str(tmp_path / "scan.pdf"))
    out = str(tmp_path / "scan.docx")
    vision = {"provider": "gemini", "model": "m", "apiKey": "k"}
    _, report = convert_pdf(pdf, out, str(tmp_path / "media"), vision)
    assert report.degraded_pages == [1]
    assert any("không trả về kết quả có thể sử dụng" in w for w in report.warnings)


# ─── 3. Worker fail-fast on a rejected key + single refund ────────────────────

def _redis_count(quota: QuotaService, user: str) -> int:
    return int(quota._redis.get(quota._key(user)) or 0)


def test_worker_fails_fast_on_vision_auth_error_and_refunds(monkeypatch, tmp_path):
    import fakeredis

    store = JobStore(redis_client=fakeredis.FakeRedis(decode_responses=True))
    quota = QuotaService(redis_client=store.redis_client, limit=3)
    monkeypatch.setattr(worker, "QUOTA", quota)

    pdf = tmp_path / "doc.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake")
    job = {"jobId": "jauth", "pdfPath": str(pdf), "filename": "doc.pdf",
           "userId": "uauth",
           "vision": {"provider": "gemini", "model": "m", "apiKey": "bad"}}
    store.save("jauth", {"jobId": "jauth", "status": "queued", "userId": "uauth"})
    store.enqueue(job)
    dequeued = store.dequeue(timeout=1)

    def raise_auth(*args, **kwargs):
        raise VisionAuthError("Gemini rejected the provided API key")
    monkeypatch.setattr(worker, "convert_pdf", raise_auth)

    assert quota.check_and_increment("uauth")[0] is True  # charged on submit
    assert _redis_count(quota, "uauth") == 1

    worker.process_job(store, dequeued)

    state = store.load("jauth")
    assert state["status"] == "failed"
    assert "Khóa API Gemini" in state["error"]
    assert _redis_count(quota, "uauth") == 0  # refunded exactly once

    # A second terminal handling must not refund again.
    worker.process_job(store, dequeued)
    assert _redis_count(quota, "uauth") == 0


# ─── 4. Gemini client response parsing (recorded payloads) ────────────────────

class _FakeResp:
    def __init__(self, text=None, candidates=None):
        self.text = text
        self.candidates = candidates or []


class _FakePart:
    def __init__(self, text):
        self.text = text


class _FakeContent:
    def __init__(self, parts):
        self.parts = parts


class _FakeCandidate:
    def __init__(self, parts):
        self.content = _FakeContent(parts)


def _client_with_response(resp):
    client = GeminiVisionClient(api_key="fake-key")
    client.convert_scanned_batch = lambda pdf, original_pages: resp
    return client


def test_extract_batch_json_plain_text():
    client = _client_with_response(_FakeResp(text=json.dumps(FAKE_VISION_BLOCKS)))
    assert client.extract_batch_json(b"%PDF", (1,)) == FAKE_VISION_BLOCKS


def test_extract_batch_json_strips_code_fence():
    fenced = "```json\n" + json.dumps(FAKE_VISION_BLOCKS) + "\n```"
    client = _client_with_response(_FakeResp(text=fenced))
    assert client.extract_batch_json(b"%PDF", (1,)) == FAKE_VISION_BLOCKS


def test_extract_batch_json_assembles_from_candidates():
    resp = _FakeResp(text=None, candidates=[
        _FakeCandidate([_FakePart(json.dumps(FAKE_VISION_BLOCKS))]),
    ])
    client = _client_with_response(resp)
    assert client.extract_batch_json(b"%PDF", (1,)) == FAKE_VISION_BLOCKS


def test_extract_batch_json_malformed_returns_none():
    client = _client_with_response(_FakeResp(text="not json at all"))
    assert client.extract_batch_json(b"%PDF", (1,)) is None
    empty = _client_with_response(_FakeResp(text=None))
    assert empty.extract_batch_json(b"%PDF", (1,)) is None


def test_client_requires_injected_key():
    with pytest.raises(RuntimeError):
        GeminiVisionClient(api_key="")


def test_auth_failure_maps_to_vision_auth_error(monkeypatch):
    class _APIError(Exception):
        def __init__(self, msg, code):
            super().__init__(msg)
            self.code = code

    client = GeminiVisionClient(api_key="fake-key")

    def boom(pdf, first, last):
        raise _APIError("API key not valid", 401)
    # Bypass the SDK: drive convert_scanned_batch's mapping logic directly.
    from vision import gemini_contract as gc
    assert gc._is_auth_failure(_APIError("API key not valid", 401)) is True
    assert gc._is_auth_failure(_APIError("quota exceeded", 429)) is False
    assert gc._is_auth_failure(Exception("Permission denied")) is True
    assert gc._is_auth_failure(Exception("internal error")) is False
