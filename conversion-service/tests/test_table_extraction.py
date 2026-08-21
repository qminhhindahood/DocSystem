"""Observable TABLE_HEAVY fidelity through a real PDF -> DOCX conversion."""
from __future__ import annotations

import fitz
from docx import Document
from docx.oxml.ns import qn

from pipeline import convert_pdf


TIMES = r"C:\Windows\Fonts\times.ttf"
TIMESBD = r"C:\Windows\Fonts\timesbd.ttf"


def _write_text(page, x: float, y: float, text: str, *, bold: bool = False) -> None:
    font = fitz.Font(fontfile=TIMESBD if bold else TIMES)
    writer = fitz.TextWriter(page.rect)
    writer.append((x, y), text, font=font, fontsize=12)
    writer.write_text(page)


def _make_table_pdf(path: str) -> str:
    doc = fitz.open()
    page = doc.new_page()
    _write_text(
        page,
        72,
        220,
        "Danh sách cán bộ tham gia công tác văn thư và lưu trữ hồ sơ",
        bold=True,
    )

    xs = [72, 230, 390, 520]
    ys = [270, 305, 340, 375]
    for x in xs:
        page.draw_line((x, ys[0]), (x, ys[-1]))
    for y in ys:
        page.draw_line((xs[0], y), (xs[-1], y))

    rows = [
        ["Họ và tên", "Chức vụ", "Đơn vị"],
        ["Nguyễn Văn An", "Chuyên viên", "Văn phòng"],
        ["Trần Thị Bình", "Kế toán", "Phòng Tài chính"],
    ]
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            _write_text(
                page,
                xs[col_index] + 6,
                ys[row_index] + 23,
                value,
                bold=row_index == 0,
            )

    page.wrap_contents()
    doc.save(path)
    doc.close()
    return path


def _make_sparse_table_pdf(path: str) -> str:
    doc = fitz.open()
    page = doc.new_page()
    _write_text(
        page,
        72,
        220,
        "Biểu mẫu lưu trữ hồ sơ hành chính của đơn vị được lập theo quy định",
    )
    xs = [72, 220, 368, 516]
    ys = [270, 305, 340, 375]
    for x in xs:
        page.draw_line((x, ys[0]), (x, ys[-1]))
    for y in ys:
        page.draw_line((xs[0], y), (xs[-1], y))
    _write_text(page, xs[0] + 6, ys[0] + 23, "Nội dung")
    _write_text(page, xs[1] + 6, ys[0] + 23, "Ghi chú")
    page.wrap_contents()
    doc.save(path)
    doc.close()
    return path


def _make_merged_header_table_pdf(path: str) -> str:
    doc = fitz.open()
    page = doc.new_page()
    _write_text(page, 72, 220, "Phụ lục danh sách cán bộ hành chính của đơn vị")
    xs = [72, 220, 368, 516]
    ys = [270, 305, 340, 375]
    page.draw_rect((xs[0], ys[0], xs[-1], ys[-1]))
    for y in ys[1:-1]:
        page.draw_line((xs[0], y), (xs[-1], y))
    for x in xs[1:-1]:
        page.draw_line((x, ys[1]), (x, ys[-1]))
    _write_text(page, xs[0] + 6, ys[0] + 23, "DANH SÁCH CÁN BỘ", bold=True)
    for column, value in enumerate(["Họ tên", "Chức vụ", "Đơn vị"]):
        _write_text(page, xs[column] + 6, ys[1] + 23, value, bold=True)
    for column, value in enumerate(["Lê Văn Cường", "Chuyên viên", "Văn phòng"]):
        _write_text(page, xs[column] + 6, ys[2] + 23, value)
    page.wrap_contents()
    doc.save(path)
    doc.close()
    return path


def _make_rowspan_table_pdf(path: str) -> str:
    doc = fitz.open()
    page = doc.new_page()
    _write_text(page, 72, 220, "Bảng phân công nhiệm vụ hành chính của đơn vị")
    xs = [72, 220, 368, 516]
    ys = [270, 305, 340, 375]
    page.draw_rect((xs[0], ys[0], xs[-1], ys[-1]))
    for x in xs[1:-1]:
        page.draw_line((x, ys[0]), (x, ys[-1]))
    page.draw_line((xs[1], ys[1]), (xs[-1], ys[1]))
    page.draw_line((xs[0], ys[2]), (xs[-1], ys[2]))
    values = [
        (0, 0, "Nhóm một"),
        (0, 1, "Nhiệm vụ A"),
        (0, 2, "Đơn vị A"),
        (1, 1, "Nhiệm vụ B"),
        (1, 2, "Đơn vị B"),
        (2, 0, "Nhóm hai"),
        (2, 1, "Nhiệm vụ C"),
        (2, 2, "Đơn vị C"),
    ]
    for row, column, value in values:
        _write_text(page, xs[column] + 6, ys[row] + 23, value)
    page.wrap_contents()
    doc.save(path)
    doc.close()
    return path


def test_table_heavy_pdf_emits_one_word_table_without_duplicate_cells(tmp_path):
    pdf = _make_table_pdf(str(tmp_path / "table.pdf"))
    out = str(tmp_path / "table.docx")

    path, report = convert_pdf(pdf, out)

    converted = Document(path)
    data_tables = [table for table in converted.tables if "Họ và tên" in table.cell(0, 0).text]
    assert len(data_tables) == 1
    assert [[cell.text for cell in row.cells] for row in data_tables[0].rows] == [
        ["Họ và tên", "Chức vụ", "Đơn vị"],
        ["Nguyễn Văn An", "Chuyên viên", "Văn phòng"],
        ["Trần Thị Bình", "Kế toán", "Phòng Tài chính"],
    ]
    all_text = " ".join(
        [paragraph.text for paragraph in converted.paragraphs]
        + [cell.text for table in converted.tables for row in table.rows for cell in row.cells]
    )
    assert all_text.count("Nguyễn Văn An") == 1
    assert all_text.count("Danh sách cán bộ tham gia") == 1
    assert report.page_types == {"TABLE_HEAVY": 1}


def test_sparse_detected_grid_uses_text_fallback_with_fidelity_warning(tmp_path):
    pdf = _make_sparse_table_pdf(str(tmp_path / "sparse.pdf"))

    path, report = convert_pdf(pdf, str(tmp_path / "sparse.docx"))

    converted = Document(path)
    text = " ".join(paragraph.text for paragraph in converted.paragraphs)
    assert "Biểu mẫu lưu trữ hồ sơ hành chính" in text
    assert report.status == "completed_with_warnings"
    assert any("failed the quality gate" in warning for warning in report.warnings)


def test_merged_pdf_cell_becomes_word_grid_span(tmp_path):
    pdf = _make_merged_header_table_pdf(str(tmp_path / "merged.pdf"))

    path, _ = convert_pdf(pdf, str(tmp_path / "merged.docx"))

    table = next(
        table for table in Document(path).tables
        if "DANH SÁCH CÁN BỘ" in table.cell(0, 0).text
    )
    physical_cells = table.rows[0]._tr.tc_lst
    assert len(physical_cells) == 1
    grid_span = physical_cells[0].tcPr.find(qn("w:gridSpan"))
    assert grid_span is not None
    assert grid_span.get(qn("w:val")) == "3"


def test_rowspan_does_not_shift_following_cells_into_merged_column(tmp_path):
    pdf = _make_rowspan_table_pdf(str(tmp_path / "rowspan.pdf"))

    path, _ = convert_pdf(pdf, str(tmp_path / "rowspan.docx"))

    table = next(
        table for table in Document(path).tables
        if "Nhóm một" in table.cell(0, 0).text
    )
    assert table.cell(0, 0).text == "Nhóm một"
    assert table.cell(1, 0).text == "Nhóm một"
    assert table.cell(1, 1).text == "Nhiệm vụ B"
    assert table.cell(1, 2).text == "Đơn vị B"
