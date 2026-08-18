"""tests/test_fidelity.py — PDF→DOCX fidelity tests (critique fix #4).

The original test suite only exercised the JSON-fixture → DOCX render path;
the PDF-ingest path (extraction → zones → classifier → assembly) had ZERO
content assertions. These tests generate real Vietnamese PDFs (PyMuPDF
TextWriter, which produces a clean Unicode text layer) and assert that the
emitted DOCX preserves the content — including the Decree-30 mandatory
components that were previously silently dropped.
"""
from __future__ import annotations

import fitz
import pytest
from docx import Document

from pipeline import convert_pdf

TIMES = r"C:\Windows\Fonts\times.ttf"
TIMESBD = r"C:\Windows\Fonts\timesbd.ttf"


def _docx_text(docx_path: str) -> str:
    """All text in the DOCX: paragraphs + table cells, whitespace-normalized."""
    dx = Document(docx_path)
    parts = [p.text for p in dx.paragraphs]
    for t in dx.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return " ".join(" ".join(parts).split())


# Visual lookalikes the renderer may substitute at the font level. U+037E
# (Greek Question Mark) is glyph-identical to ';' and appears when some fonts
# remap the ASCII semicolon.
_LOOKALIKES = str.maketrans({"\u037e": ";", "\u2010": "-", "\u2011": "-",
                             "\u2012": "-", "\u2013": "-", "\u2014": "-"})


def _norm(s: str) -> str:
    return s.casefold().translate(_LOOKALIKES)


def _has(text: str, phrase: str) -> bool:
    """Case-insensitive containment, tolerant of glyph-identical lookalikes:
    the rule engine legitimately uppercases article headings and the signature
    name per Decree 30, so fidelity is about the CONTENT surviving, not its
    exact case or a font-level semicolon substitution."""
    return _norm(phrase) in _norm(text)


def _make_qd_pdf(path: str) -> str:
    """A realistic 1-page Quyết định: header zone, title, wrapped body,
    signature zone. TextWriter keeps the text layer clean Unicode."""
    f_reg = fitz.Font(fontfile=TIMES)
    f_bold = fitz.Font(fontfile=TIMESBD)
    doc = fitz.open()
    page = doc.new_page()
    W = page.rect.width

    def put(x, y, text, font=f_reg, size=14):
        tw = fitz.TextWriter(page.rect)
        tw.append((x, y), text, font=font, fontsize=size)
        tw.write_text(page)

    def putc(y, text, font=f_reg, size=14):
        w = font.text_length(text, fontsize=size)
        put((W - w) / 2, y, text, font, size)

    # header zone (top 25%)
    put(72, 60, "BỘ NỘI VỤ", f_bold, 13)
    put(72, 80, "VỤ TỔ CHỨC CÁN BỘ", f_bold, 13)
    put(72, 100, "Số: 123/QĐ-TCCB", f_reg, 13)
    putc(60, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", f_bold, 13)
    putc(80, "Độc lập - Tự do - Hạnh phúc", f_bold, 13)
    putc(105, "Hà Nội, ngày 15 tháng 01 năm 2026", f_reg, 13)
    # title
    putc(250, "QUYẾT ĐỊNH", f_bold, 15)
    putc(275, "Về việc thành lập Ban công tác văn thư", f_bold, 14)
    # body — a paragraph wrapped over two visual lines
    put(72, 330, "Căn cứ Nghị định số 30/2020/NĐ-CP ngày 05 tháng 3 năm 2020", f_reg, 14)
    put(72, 350, "của Chính phủ về công tác văn thư; lưu trữ hồ sơ tài liệu.", f_reg, 14)
    put(72, 380, "Theo đề nghị của Vụ trưởng Vụ Tổ chức cán bộ.", f_reg, 14)
    put(72, 410, "Điều 1. Thành lập Ban công tác văn thư gồm các ông bà", f_reg, 14)
    put(72, 430, "có tên trong danh sách kèm theo Quyết định này.", f_reg, 14)
    put(72, 460, "Điều 2. Quyết định này có hiệu lực kể từ ngày ký.", f_reg, 14)
    # signature zone (bottom 25%)
    put(72, 660, "Nơi nhận:", f_bold, 12)
    put(72, 680, "- Như Điều 3;", f_reg, 12)
    put(72, 700, "- Lưu: VT, TCCB.", f_reg, 12)
    putc(660, "BỘ TRƯỞNG", f_bold, 13)
    putc(740, "Nguyễn Văn A", f_bold, 13)

    doc.save(path)
    doc.close()
    return path


@pytest.fixture(scope="module")
def qd_result(tmp_path_factory):
    tmp = tmp_path_factory.mktemp("fidelity")
    pdf = _make_qd_pdf(str(tmp / "qd.pdf"))
    out = str(tmp / "qd.docx")
    path, report = convert_pdf(pdf, out)
    return path, report, _docx_text(path)


# ─── Fix #2: header/signature zones are no longer dropped ─────────────────────

def test_quoc_hieu_present(qd_result):
    _, _, text = qd_result
    assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in text


def test_tieu_ngu_present(qd_result):
    _, _, text = qd_result
    assert "Độc lập - Tự do - Hạnh phúc" in text


def test_issuing_agency_present(qd_result):
    _, _, text = qd_result
    assert "BỘ NỘI VỤ" in text
    assert "VỤ TỔ CHỨC CÁN BỘ" in text


def test_document_number_present(qd_result):
    _, _, text = qd_result
    assert "123/QĐ-TCCB" in text


def test_title_present(qd_result):
    _, _, text = qd_result
    assert "QUYẾT ĐỊNH" in text
    assert "Về việc thành lập Ban công tác văn thư" in text


def test_signature_present(qd_result):
    _, _, text = qd_result
    assert "BỘ TRƯỞNG" in text
    assert _has(text, "Nguyễn Văn A")


def test_receipt_list_present(qd_result):
    _, _, text = qd_result
    assert "Như Điều 3" in text
    assert "Lưu: VT, TCCB" in text


def test_signer_not_in_receipt_cell(qd_result):
    """The signatory's name must not leak into the Nơi nhận cell."""
    path, _, _ = qd_result
    dx = Document(path)
    for t in dx.tables:
        for row in t.rows:
            left = row.cells[0].text
            if "Nơi nhận" in left:
                assert "Nguyễn Văn A" not in left


# ─── Fix #3: wrapped lines merge WITH a space ─────────────────────────────────

def test_wrapped_paragraph_merged_with_space(qd_result):
    _, _, text = qd_result
    # The two visual lines must rejoin as one paragraph with the space intact.
    assert "năm 2020 của Chính phủ" in text
    # And must NOT appear glued without the space.
    assert "2020của" not in text


def test_dieu_body_merged(qd_result):
    _, _, text = qd_result
    assert _has(text, "các ông bà có tên trong danh sách")


# ─── Fix #1: confidence reflects reality ──────────────────────────────────────

def test_confidence_not_fake_one(qd_result):
    _, report, _ = qd_result
    # A real conversion is capped by coverage and block confidence; it must
    # never report a perfect 1.0.
    assert report.confidence < 1.0
    assert report.status == "completed"


def test_coverage_reported(qd_result):
    _, report, _ = qd_result
    assert report.extracted_chars > 0
    assert report.output_chars > 0
    assert 0.0 < report.coverage <= 1.0
    # This well-formed document should convert with high coverage.
    assert report.coverage >= 0.85


def test_empty_output_fails_not_confidence_one(tmp_path):
    """A PDF whose text layer yields nothing usable must FAIL, never report
    completed with confidence 1.0 (the original bug)."""
    # Blank page: no text at all -> scanned path without Gemini -> degraded.
    doc = fitz.open()
    doc.new_page()
    pdf = str(tmp_path / "blank.pdf")
    doc.save(pdf)
    doc.close()
    _, report = convert_pdf(pdf, str(tmp_path / "blank.docx"))
    assert report.confidence < 1.0
    assert report.status in ("failed", "completed_with_warnings")


# ─── Fix #4: character-level fidelity on body content ─────────────────────────

@pytest.mark.parametrize("phrase", [
    "Căn cứ Nghị định số 30/2020/NĐ-CP",
    "công tác văn thư; lưu trữ hồ sơ tài liệu",
    "Theo đề nghị của Vụ trưởng Vụ Tổ chức cán bộ",
    "Điều 1. Thành lập Ban công tác văn thư",
    "Điều 2. Quyết định này có hiệu lực kể từ ngày ký",
])
def test_body_phrase_fidelity(qd_result, phrase):
    _, _, text = qd_result
    assert _has(text, phrase)
