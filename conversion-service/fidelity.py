"""fidelity.py — per-job fidelity ledger (ticket 05).

Tier-1 guarantee: every extracted character lands verbatim in the DOCX,
modulo DOCUMENTED normalization:

1. case-fold — Decree-30 typography uppercases article openers and
   structural labels by design ("Điều 1." -> "ĐIỀU 1."), so case-only
   drift is invisible to this ledger (stated limitation, never silent:
   the ledger reports the normalization it applied);
2. whitespace collapse — visual line wrapping and zone grouping are
   layout artifacts, not content;
3. block order — the guarantee is per-character presence (bag compare),
   not document order: header/signature grouping reorders blocks.

Bag (multiset) fidelity is immune to reordering and the uppercase
transform and exact on every other character. The reported CER is the
bag divergence rate: (chars missing from DOCX + chars extra in DOCX) /
max(len(extracted), len(docx)) — a substitution counts as 2 units, so
the ledger is never optimistic (bag-CER >= Levenshtein-CER for single
edits).

Scanned pages are transcribed (OCR), never copied — for jobs with any
scanned page the ledger is explicitly None: no fake fidelity number.
"""
from __future__ import annotations

from collections import Counter
from dataclasses import dataclass, field

MAX_SPANS = 10  # divergence spans are for human review — capped payload


@dataclass(frozen=True)
class FidelityLedger:
    """Bag fidelity of DOCX output vs extracted text (documented norms)."""

    fidelity: float
    cer: float                      # bag divergence rate (see module docstring)
    extracted_chars: int
    docx_chars: int
    divergence_spans: list[dict] = field(default_factory=list)

    MAX_SPANS: int = MAX_SPANS

    def to_payload(self) -> dict:
        """JSON-safe dict for the job report payload."""
        return {
            "fidelity": round(self.fidelity, 4),
            "cer": round(self.cer, 4),
            "extractedChars": self.extracted_chars,
            "docxChars": self.docx_chars,
            "divergenceSpans": self.divergence_spans,
            "normalization": ["casefold", "whitespace-collapse", "bag-order-free"],
        }


def _normalize(text: str) -> str:
    """Documented normalization: case-fold + whitespace collapse."""
    return " ".join(text.casefold().split())


def _docx_flat_text(docx_path: str) -> str:
    """Flatten a DOCX to comparable text: paragraphs + table cells."""
    from docx import Document

    d = Document(docx_path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return " ".join(parts)


def _divergence_spans(ext_words: Counter, docx_words: Counter) -> list[dict]:
    """Word-level divergence spans (missing + extra), capped at MAX_SPANS."""
    spans: list[dict] = []
    missing = ext_words - docx_words
    extra = docx_words - ext_words
    for word, count in sorted(missing.items(), key=lambda kv: -kv[1]):
        if len(spans) >= MAX_SPANS:
            break
        spans.append({"kind": "missing", "extracted": word, "docx": "", "count": count})
    for word, count in sorted(extra.items(), key=lambda kv: -kv[1]):
        if len(spans) >= MAX_SPANS:
            break
        spans.append({"kind": "extra", "extracted": "", "docx": word, "count": count})
    return spans


def compute_ledger(extracted_text: str, docx_text: str) -> FidelityLedger:
    """Compare extracted text vs DOCX text under the documented norms."""
    ext = _normalize(extracted_text)
    doc = _normalize(docx_text)
    ext_chars, docx_chars = Counter(ext), Counter(doc)
    missing = sum((ext_chars - docx_chars).values())
    extra = sum((docx_chars - ext_chars).values())
    denom = max(len(ext), len(doc), 1)
    cer = (missing + extra) / denom
    return FidelityLedger(
        fidelity=max(0.0, 1.0 - cer),  # clamp: catastrophic divergence reports 0, never negative
        cer=cer,
        extracted_chars=len(ext),
        docx_chars=len(doc),
        divergence_spans=_divergence_spans(Counter(ext.split()), Counter(doc.split())),
    )


def ledger_for_docx(extracted_text: str, docx_path: str) -> FidelityLedger:
    """Convenience: compute the ledger against a saved DOCX file."""
    return compute_ledger(extracted_text, _docx_flat_text(docx_path))
