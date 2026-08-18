"""render/docx_builder.py — DocxBlockBuilder (python-docx).

Consumes the validated JSON contract and emits DOCX with Decree 30
typography, sections + landscape orientation, and the manual wp:anchor
floating-image helper for seals (plan §7 render note). python-docx has NO
high-level API for floating images, so the wp:anchor element is built by hand.
"""
from __future__ import annotations

import itertools
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt

from rules.rule_engine import RuleEngine
from schema.blocks import (
    AdminHeaderBlock,
    Block,
    HeadingBlock,
    IllegibleBlock,
    ImageBlock,
    ListBlock,
    ParagraphBlock,
    Run,
    SectionBreakBlock,
    SignatureBlock,
    TableBlock,
    TableCell,
)

# A4 in EMU (1 mm = 36000 EMU). Used to convert the 0-1000 bbox grid.
A4_WIDTH_EMU = 210 * 36000   # 7,560,000
A4_HEIGHT_EMU = 297 * 36000  # 10,692,000

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_drawing_id = itertools.count(1000)


def _next_drawing_id() -> int:
    return next(_drawing_id)


# ─── Floating image (wp:anchor) — plan §7 render note ─────────────────────────

def _build_pic_graphic(rId: str, width_emu: int, height_emu: int) -> OxmlElement:
    """a:graphic > a:graphicData > pic:pic (nvPicPr + blipFill + spPr/a:xfrm)."""
    graphic = OxmlElement("a:graphic")
    graphic_data = OxmlElement("a:graphicData")
    graphic_data.set(
        "uri", "http://schemas.openxmlformats.org/drawingml/2006/picture"
    )

    pic = OxmlElement("pic:pic")

    nv = OxmlElement("pic:nvPicPr")
    cnv_pr = OxmlElement("pic:cNvPr")
    cnv_pr.set("id", str(_next_drawing_id()))
    cnv_pr.set("name", "seal")
    nv.append(cnv_pr)
    nv.append(OxmlElement("pic:cNvPicPr"))
    pic.append(nv)

    blip_fill = OxmlElement("pic:blipFill")
    blip = OxmlElement("a:blip")
    blip.set(qn("r:embed"), rId)
    blip_fill.append(blip)
    stretch = OxmlElement("a:stretch")
    stretch.append(OxmlElement("a:fillRect"))
    blip_fill.append(stretch)
    pic.append(blip_fill)

    sp_pr = OxmlElement("pic:spPr")
    xfrm = OxmlElement("a:xfrm")
    off = OxmlElement("a:off")
    off.set("x", "0")
    off.set("y", "0")
    xfrm.append(off)
    ext = OxmlElement("a:ext")
    ext.set("cx", str(int(width_emu)))
    ext.set("cy", str(int(height_emu)))
    xfrm.append(ext)
    sp_pr.append(xfrm)
    prst = OxmlElement("a:prstGeom")
    prst.set("prst", "rect")
    prst.append(OxmlElement("a:avLst"))
    sp_pr.append(prst)
    pic.append(sp_pr)

    graphic_data.append(pic)
    graphic.append(graphic_data)
    return graphic


def add_floating_image(
    paragraph: Any,
    image_path: str | Path,
    left_emu: int,
    top_emu: int,
    width_emu: int,
    height_emu: int,
) -> None:
    """Build a wp:anchor floating image manually (no python-docx API exists).

    Seals render ON TOP of text (behindDoc='0'), like a real stamp over the
    signature name. Position is absolute relative to the PAGE, converted from
    the PDF bbox: offset_emu = bbox/1000 * page_dimension_emu.
    """
    run = paragraph.add_run()
    # python-docx 1.2: DocumentPart.get_or_add_image() -> (rId, Image)
    rId, _image = paragraph.part.get_or_add_image(str(image_path))

    anchor = OxmlElement("wp:anchor")
    anchor.set("behindDoc", "0")        # on top of text, like a real stamp
    anchor.set("locked", "1")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "251659264")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")

    for axis, offset in (("positionH", left_emu), ("positionV", top_emu)):
        pos = OxmlElement(f"wp:{axis}")
        pos.set("relativeFrom", "page")
        pos_off = OxmlElement("wp:posOffset")
        pos_off.text = str(int(offset))
        pos.append(pos_off)
        anchor.append(pos)

    extent = OxmlElement("wp:extent")
    extent.set("cx", str(int(width_emu)))
    extent.set("cy", str(int(height_emu)))
    anchor.append(extent)

    anchor.append(OxmlElement("wp:effectExtent"))
    anchor.append(OxmlElement("wp:wrapNone"))  # seal overlaps text — no wrap

    doc_pr = OxmlElement("wp:docPr")
    doc_pr.set("id", str(_next_drawing_id()))
    doc_pr.set("name", "seal")
    anchor.append(doc_pr)

    anchor.append(_build_pic_graphic(rId, width_emu, height_emu))

    drawing = OxmlElement("w:drawing")
    drawing.append(anchor)
    run._r.append(drawing)


def bbox_to_emu(bbox: list[int]) -> tuple[int, int, int, int]:
    """Convert a 0-1000 grid bbox [x0,y0,x1,y1] to absolute page EMU."""
    x0, y0, x1, y1 = bbox
    left = int(x0 / 1000 * A4_WIDTH_EMU)
    top = int(y0 / 1000 * A4_HEIGHT_EMU)
    width = int((x1 - x0) / 1000 * A4_WIDTH_EMU)
    height = int((y1 - y0) / 1000 * A4_HEIGHT_EMU)
    return left, top, width, height


# ─── The builder ──────────────────────────────────────────────────────────────

class DocxBlockBuilder:
    """Render a validated block list to a DOCX Document."""

    def __init__(self, rule_engine: RuleEngine):
        self.rules = rule_engine
        self.doc = Document()
        self._configure_base_section()

    # ── setup ────────────────────────────────────────────────────────────────
    def _configure_base_section(self) -> None:
        section = self.doc.sections[0]
        m = self.rules.margins_mm
        section.top_margin = Mm(m["top"])
        section.bottom_margin = Mm(m["bottom"])
        section.left_margin = Mm(m["left"])
        section.right_margin = Mm(m["right"])
        section.page_width = Mm(210)
        section.page_height = Mm(297)

    def _set_normal_style(self) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = self.rules.font_family
        style.font.size = Pt(self.rules.body_size_pt)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(attr), self.rules.font_family)

    # ── public ───────────────────────────────────────────────────────────────
    def build(self, blocks: list[Block]) -> Document:
        self._set_normal_style()
        for block in blocks:
            self._render_block(block)
        return self.doc

    def save(self, blocks: list[Block], out_path: str | Path) -> Path:
        doc = self.build(blocks)
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return out

    # ── dispatch ─────────────────────────────────────────────────────────────
    def _render_block(self, block: Block) -> None:
        if isinstance(block, AdminHeaderBlock):
            self._render_admin_header(block)
        elif isinstance(block, HeadingBlock):
            self._render_heading(block)
        elif isinstance(block, ParagraphBlock):
            self._render_paragraph(block)
        elif isinstance(block, ListBlock):
            self._render_list(block)
        elif isinstance(block, TableBlock):
            self._render_table(block)
        elif isinstance(block, ImageBlock):
            self._render_image(block)
        elif isinstance(block, SignatureBlock):
            self._render_signature(block)
        elif isinstance(block, SectionBreakBlock):
            self._render_section_break(block)
        elif isinstance(block, IllegibleBlock):
            self._render_illegible(block)

    # ── run helpers ──────────────────────────────────────────────────────────
    def _add_runs(
        self,
        paragraph: Any,
        runs: list[Run],
        *,
        size_pt: Optional[float] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        upper: bool = False,
    ) -> None:
        for r in runs:
            run = paragraph.add_run(r.text.upper() if upper else r.text)
            run.font.name = self.rules.font_family
            if size_pt:
                run.font.size = Pt(size_pt)
            run.font.bold = r.bold if r.bold is not None else bold
            run.font.italic = r.italic if r.italic is not None else italic
            if r.underline:
                run.font.underline = True

    def _add_text(
        self,
        paragraph: Any,
        text: str,
        *,
        size_pt: Optional[float] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        upper: bool = False,
        underline: bool = False,
    ) -> None:
        run = paragraph.add_run(text.upper() if upper else text)
        run.font.name = self.rules.font_family
        if size_pt:
            run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        if underline:
            run.font.underline = True

    # ── block renderers ──────────────────────────────────────────────────────
    def _render_admin_header(self, block: AdminHeaderBlock) -> None:
        style = self.rules.block_style("admin_header")
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._strip_borders(table)

        left_cell, right_cell = table.rows[0].cells
        self._fill_admin_side(left_cell, block.left, style.get("left", {}))
        self._fill_admin_side(right_cell, block.right, style.get("right", {}))

    def _fill_admin_side(self, cell: Any, data: Any, style: dict[str, Any]) -> None:
        # clear the default empty paragraph
        cell.paragraphs[0].text = ""
        first = True
        for field_name in (
            "superior_agency", "issuing_agency", "document_number",
            "country_name", "motto", "location_and_date",
        ):
            value = getattr(data, field_name, None)
            if not value:
                continue
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            fs = style.get(field_name, {})
            p.alignment = _ALIGN.get(fs.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
            self._add_text(
                p, value,
                size_pt=fs.get("size_pt"),
                bold=fs.get("bold"),
                italic=fs.get("italic"),
                upper=fs.get("case") == "upper",
                underline=fs.get("underline") == "motto",
            )

    def _render_heading(self, block: HeadingBlock) -> None:
        style = self.rules.block_style("heading")
        key = f"level_{min(block.level, 3)}"
        fs = style.get(key, {})
        p = self.doc.add_paragraph()
        p.alignment = _ALIGN.get(
            block.align or fs.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT
        )
        self._add_text(
            p, block.text,
            size_pt=fs.get("size_pt", self.rules.body_size_pt),
            bold=fs.get("bold", True),
            upper=fs.get("case") == "upper",
        )

    def _render_paragraph(self, block: ParagraphBlock) -> None:
        style = self.rules.block_style("paragraph")
        p = self.doc.add_paragraph()
        p.alignment = _ALIGN.get(
            block.align or style.get("align", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        pf = p.paragraph_format
        pf.line_spacing = self.rules.line_spacing
        indent = block.first_line_indent_pt
        if indent:
            pf.first_line_indent = Pt(indent)
        runs = block.runs or ([Run(text=block.text)] if block.text else [])
        self._add_runs(p, runs, size_pt=style.get("size_pt"))

    def _render_list(self, block: ListBlock) -> None:
        style = self.rules.block_style("list")
        for idx, item in enumerate(block.items):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing = self.rules.line_spacing
            marker = self._marker_for(block, idx)
            runs = item.runs or [Run(text=item.text)]
            self._add_text(p, marker + " ", size_pt=style.get("size_pt"))
            self._add_runs(p, runs, size_pt=style.get("size_pt"))

    @staticmethod
    def _marker_for(block: ListBlock, idx: int) -> str:
        if not block.ordered:
            return block.marker or "-"
        marker = block.marker or "a)"
        if marker.endswith(")") and len(marker) >= 2 and marker[0].isalpha():
            return chr(ord(marker[0]) + idx) + ")"
        if marker.endswith(".") and marker[:-1].isdigit():
            return f"{int(marker[:-1]) + idx}."
        return marker

    def _render_table(self, block: TableBlock) -> None:
        style = self.rules.block_style("table")
        all_rows = block.headers + block.rows
        if not all_rows:
            return
        n_cols = max(len(r) for r in all_rows)
        table = self.doc.add_table(rows=len(all_rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for r_idx, row in enumerate(all_rows):
            is_header = r_idx < len(block.headers)
            col_cursor = 0
            for cell in row:
                if col_cursor >= n_cols:
                    break
                dc = table.cell(r_idx, col_cursor)
                # apply colspan/rowspan via merge
                end_col = min(col_cursor + cell.colspan, n_cols) - 1
                end_row = min(r_idx + cell.rowspan, len(all_rows)) - 1
                if end_col > col_cursor or end_row > r_idx:
                    dc = dc.merge(table.cell(end_row, end_col))
                p = dc.paragraphs[0]
                p.alignment = _ALIGN.get(cell.align or "left", WD_ALIGN_PARAGRAPH.LEFT)
                bold = cell.bold if cell.bold is not None else (
                    style.get("header_bold", True) if is_header else None
                )
                self._add_text(
                    p, cell.text,
                    size_pt=style.get("size_pt"),
                    bold=bold,
                )
                col_cursor += cell.colspan

    def _render_image(self, block: ImageBlock) -> None:
        if not block.src:
            return
        src = Path(block.src)
        if not src.exists():
            return
        if block.placement == "floating":
            p = self.doc.add_paragraph()
            left, top, width, height = bbox_to_emu(block.bbox)
            add_floating_image(p, src, left, top, width, height)
        else:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            _, _, width, height = bbox_to_emu(block.bbox)
            run.add_picture(str(src), width=Emu(width), height=Emu(height))

    def _render_signature(self, block: SignatureBlock) -> None:
        style = self.rules.block_style("signature")
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._strip_borders(table)
        left_cell, right_cell = table.rows[0].cells
        self._fill_signature_left(left_cell, block, style.get("left", {}))
        self._fill_signature_right(right_cell, block, style.get("right", {}))

    def _fill_signature_left(self, cell: Any, block: SignatureBlock, style: dict) -> None:
        cell.paragraphs[0].text = ""
        label = style.get("label", {})
        p = cell.paragraphs[0]
        self._add_text(
            p, label.get("text", "Nơi nhận:"),
            size_pt=label.get("size_pt"),
            bold=label.get("bold", True),
            italic=label.get("italic", True),
        )
        item_style = style.get("receipt_item", {})
        for item in block.left.receipt_list:
            ip = cell.add_paragraph()
            self._add_text(ip, item, size_pt=item_style.get("size_pt"))

    def _fill_signature_right(self, cell: Any, block: SignatureBlock, style: dict) -> None:
        cell.paragraphs[0].text = ""
        first = True

        def para() -> Any:
            nonlocal first
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            return p

        auth = style.get("authority", {})
        if block.right.authority:
            p = para()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_text(
                p, block.right.authority, size_pt=auth.get("size_pt"),
                bold=auth.get("bold", True), upper=auth.get("case") == "upper",
            )
        title = style.get("title", {})
        if block.right.title:
            p = para()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_text(
                p, block.right.title, size_pt=title.get("size_pt"),
                bold=title.get("bold", True), upper=title.get("case") == "upper",
            )
        # signature space
        space_lines = style.get("signature_space_lines", 3)
        for _ in range(space_lines):
            para()
        name = style.get("name", {})
        if block.right.name:
            p = para()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_text(
                p, block.right.name, size_pt=name.get("size_pt"),
                bold=name.get("bold", True), upper=name.get("case") == "upper",
            )

    def _render_section_break(self, block: SectionBreakBlock) -> None:
        section = self.doc.add_section(WD_SECTION_START.NEW_PAGE)
        if block.orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            # swap dimensions for landscape
            section.page_width, section.page_height = Mm(297), Mm(210)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width, section.page_height = Mm(210), Mm(297)
        m = self.rules.margins_mm
        section.top_margin = Mm(m["top"])
        section.bottom_margin = Mm(m["bottom"])
        section.left_margin = Mm(m["left"])
        section.right_margin = Mm(m["right"])
        if block.label:
            style = self.rules.block_style("section_break").get("label", {})
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_text(
                p, block.label, size_pt=style.get("size_pt"),
                bold=style.get("bold", True), upper=style.get("case") == "upper",
            )

    def _render_illegible(self, block: IllegibleBlock) -> None:
        style = self.rules.block_style("illegible")
        template = style.get("marker_template", "[Vùng không đọc được]")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_text(
            p, template.format(page=block.page or "?"),
            italic=style.get("italic", True),
        )

    # ── table border helper ──────────────────────────────────────────────────
    @staticmethod
    def _strip_borders(table: Any) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "none")
            borders.append(el)
        tbl_pr.append(borders)
