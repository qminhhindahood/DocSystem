"""eval/verify_p0a.py — P0a render-verification gate.

Loads each fixture, renders DOCX via DocxBlockBuilder, then RE-OPENS the DOCX
with python-docx and asserts:
  1. Decree 30 checklist: Times New Roman, A4 margins (25/25/30/20 mm),
     2-column admin header, motto underline, landscape annex section.
  2. Seal fixture: a wp:anchor element exists with behindDoc='0', wrapNone,
     and page-relative position offsets matching the fixture bbox.

Exit code 0 = gate passed.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

SERVICE_ROOT = Path(__file__).resolve().parent.parent
if str(SERVICE_ROOT) not in sys.path:
    sys.path.insert(0, str(SERVICE_ROOT))

from docx import Document  # noqa: E402
from docx.enum.section import WD_ORIENT  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Mm  # noqa: E402

import config  # noqa: E402
from render.docx_builder import (  # noqa: E402
    A4_HEIGHT_EMU,
    A4_WIDTH_EMU,
    DocxBlockBuilder,
    bbox_to_emu,
)
from rules.rule_engine import RuleEngine  # noqa: E402
from schema.blocks import ImageBlock, parse_blocks  # noqa: E402

FIXTURES = ["quyet_dinh.json", "cong_van.json", "thong_bao.json"]
OUT_DIR = SERVICE_ROOT / "eval" / "out"

FAILURES: list[str] = []


def check(name: str, cond: bool, detail: str = "") -> None:
    status = "PASS" if cond else "FAIL"
    print(f"  [{status}] {name}" + (f" — {detail}" if detail and not cond else ""))
    if not cond:
        FAILURES.append(f"{name}: {detail}")


def find_anchors(doc: Document) -> list:
    anchors = []
    for p in doc.paragraphs:
        for drawing in p._p.iter(qn("w:drawing")):
            anchors.extend(drawing.iter(qn("wp:anchor")))
    return anchors


def verify_fixture(fixture_name: str) -> None:
    print(f"\n=== {fixture_name} ===")
    fixture_path = SERVICE_ROOT / "eval" / "fixtures" / fixture_name
    data = json.loads(fixture_path.read_text(encoding="utf-8"))
    blocks = parse_blocks(data)

    # Resolve fixture-relative image src paths against the service root.
    for b in blocks:
        if isinstance(b, ImageBlock) and b.src and not Path(b.src).is_absolute():
            b.src = str(SERVICE_ROOT / b.src)

    rules = RuleEngine(config.SHARED_TYPOGRAPHY_PATH)
    rules.apply(blocks)

    out_path = OUT_DIR / fixture_name.replace(".json", ".docx")
    builder = DocxBlockBuilder(rules)
    builder.save(blocks, out_path)
    print(f"  rendered -> {out_path}")

    # Re-open and inspect
    doc = Document(str(out_path))

    # 1. Margins (first section)
    s0 = doc.sections[0]
    check("margin top 25mm", abs(s0.top_margin.mm - 25) < 0.5, f"{s0.top_margin.mm}")
    check("margin bottom 25mm", abs(s0.bottom_margin.mm - 25) < 0.5, f"{s0.bottom_margin.mm}")
    check("margin left 30mm", abs(s0.left_margin.mm - 30) < 0.5, f"{s0.left_margin.mm}")
    check("margin right 20mm", abs(s0.right_margin.mm - 20) < 0.5, f"{s0.right_margin.mm}")

    # 2. Font: Normal style is Times New Roman
    normal = doc.styles["Normal"]
    check("font Times New Roman", normal.font.name == "Times New Roman", str(normal.font.name))

    # 3. Landscape annex section exists
    landscape = [s for s in doc.sections if s.orientation == WD_ORIENT.LANDSCAPE]
    check("landscape annex section", len(landscape) >= 1, f"{len(landscape)} landscape sections")
    if landscape:
        ls = landscape[0]
        check(
            "landscape dims swapped",
            ls.page_width.mm > ls.page_height.mm,
            f"{ls.page_width.mm}x{ls.page_height.mm}",
        )

    # 4. 2-column admin header: a borderless 2-col table near the top
    tables = doc.tables
    check("has tables (header/signature/annex)", len(tables) >= 2, f"{len(tables)} tables")
    if tables:
        check("admin header 2 columns", len(tables[0].columns) == 2,
              f"{len(tables[0].columns)} cols")

    # 5. Motto underline present somewhere
    underlined = False
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.underline:
                underlined = True
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.font.underline:
                            underlined = True
    check("motto underline present", underlined)

    # 6. Floating seal: wp:anchor with behindDoc='0', wrapNone, page offsets
    seals = [b for b in blocks if isinstance(b, ImageBlock) and b.kind == "seal"]
    anchors = find_anchors(doc)
    check("wp:anchor present", len(anchors) >= len(seals),
          f"{len(anchors)} anchors vs {len(seals)} seals")
    for anchor in anchors:
        check("anchor behindDoc='0'", anchor.get("behindDoc") == "0",
              str(anchor.get("behindDoc")))
        wrap_none = anchor.find(qn("wp:wrapNone"))
        check("anchor wrapNone", wrap_none is not None)
        pos_h = anchor.find(qn("wp:positionH"))
        pos_v = anchor.find(qn("wp:positionV"))
        check("anchor positionH relativeFrom=page",
              pos_h is not None and pos_h.get("relativeFrom") == "page")
        check("anchor positionV relativeFrom=page",
              pos_v is not None and pos_v.get("relativeFrom") == "page")

    # 7. Offsets match the fixture bbox (EMU conversion)
    if seals and anchors:
        seal = seals[0]
        exp_left, exp_top, exp_w, exp_h = bbox_to_emu(seal.bbox)
        anchor = anchors[0]
        got_left = int(anchor.find(qn("wp:positionH")).find(qn("wp:posOffset")).text)
        got_top = int(anchor.find(qn("wp:positionV")).find(qn("wp:posOffset")).text)
        extent = anchor.find(qn("wp:extent"))
        got_w = int(extent.get("cx"))
        got_h = int(extent.get("cy"))
        check("anchor left offset", abs(got_left - exp_left) <= 1, f"{got_left} vs {exp_left}")
        check("anchor top offset", abs(got_top - exp_top) <= 1, f"{got_top} vs {exp_top}")
        check("anchor width", abs(got_w - exp_w) <= 1, f"{got_w} vs {exp_w}")
        check("anchor height", abs(got_h - exp_h) <= 1, f"{got_h} vs {exp_h}")


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for f in FIXTURES:
        verify_fixture(f)
    print()
    if FAILURES:
        print(f"P0a GATE: FAIL ({len(FAILURES)} failures)")
        for f in FAILURES:
            print("  -", f)
        return 1
    print("P0a GATE: PASS — all fixtures render with Decree 30 structure + floating seals")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
