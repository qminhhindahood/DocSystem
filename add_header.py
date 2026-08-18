"""
Script: Thêm header chuẩn văn bản hành chính cho các file templates
Header: Bộ GD&ĐT / CHXHCNVN / Độc lập - Tự do - Hạnh phúc

Sử dụng:
    python add_header.py              # Xử lý thư mục templates/
    python add_header.py --templates-dir path/to/templates
    python add_header.py --dry-run    # Chỉ xem thay đổi, không ghi file
    python add_header.py --backup     # Tạo backup trước khi ghi
"""

import os
import argparse
import shutil
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import lxml.etree as etree

# --- Cấu hình font chuẩn ---
FONT_NAME = "Times New Roman"
FONT_SIZE_HEADER = Pt(13)      # Size cho dòng Bộ GD&ĐT / CHXHCNVN
FONT_SIZE_HEADER_SMALL = Pt(12)  # Size cho "Độc lập - Tự do - Hạnh phúc"
FONT_SIZE_META = Pt(13)         # Size cho dòng Số: / Thành phố
FONT_SIZE_TITLE = Pt(14)        # Size cho tiêu đề THÔNG TƯ

# --- Text chuẩn ---
HEADER_LEFT = "BỘ GIÁO DỤC VÀ ĐÀO TẠO"
HEADER_CENTER = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
HEADER_RIGHT = "Độc lập - Tự do - Hạnh phúc"


def set_cell_width(cell, width_cm):
    """Đặt độ rộng cố định cho ô"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))  # 1 cm ≈ 567 twips
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_font(run, size=None, bold=False, italic=False, color=None):
    """Áp dụng font chữ cho một run"""
    font = run.font
    font.name = FONT_NAME
    font.size = size or Pt(13)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
    # Áp dụng font cho East Asian text
    r = run._r
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def set_paragraph_spacing(para, before=0, after=0, line=None):
    """Đặt khoảng cách đoạn"""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")


def set_paragraph_format(para, alignment=None, space_before=0, space_after=0, line_spacing=None):
    """Định dạng đoạn văn"""
    para.alignment = alignment
    set_paragraph_spacing(para, space_before, space_after, line_spacing)


def add_formatted_run(para, text, size=None, bold=False, italic=False, color=None, font_name=None):
    """Thêm text đã format vào đoạn"""
    run = para.add_run(text)
    f = run.font
    f.name = font_name or FONT_NAME
    f.size = size or Pt(13)
    f.bold = bold
    f.italic = italic
    if color:
        f.color.rgb = color
    # East Asian font
    r = run._r
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name or FONT_NAME)
    rFonts.set(qn("w:hAnsi"), font_name or FONT_NAME)
    rFonts.set(qn("w:cs"), font_name or FONT_NAME)
    rFonts.set(qn("w:eastAsia"), font_name or FONT_NAME)
    return run


def build_header_table(doc):
    """Tạo bảng header 1x3: Bộ | CHXHCNVN | Độc lập"""

    # Xóa tất cả paragraph trống ở đầu document (trước khi thêm table)
    # Giữ lại các paragraph có content
    body = doc.element.body
    # Chèn table sau các element đầu tiên (nếu có), hoặc append
    table = doc.add_table(rows=1, cols=3)
    table.allow_autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tắt border của bảng
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Cell trái - BỘ GIÁO DỤC VÀ ĐÀO TẠO
    cell_left = table.cell(0, 0)
    cell_left.width = Cm(5.5)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(p_left, space_before=0, space_after=60, line_spacing=240)
    add_formatted_run(p_left, HEADER_LEFT, size=FONT_SIZE_HEADER, bold=True)

    # Cell giữa - CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
    cell_center = table.cell(0, 1)
    cell_center.width = Cm(7)
    p_center = cell_center.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p_center, space_before=0, space_after=60, line_spacing=240)
    add_formatted_run(p_center, HEADER_CENTER, size=FONT_SIZE_HEADER, bold=True)

    # Cell phải - Độc lập - Tự do - Hạnh phúc
    cell_right = table.cell(0, 2)
    cell_right.width = Cm(5.5)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(p_right, space_before=0, space_after=60, line_spacing=240)
    add_formatted_run(p_right, HEADER_RIGHT, size=FONT_SIZE_HEADER_SMALL, bold=False)

    return table


def add_meta_line(doc):
    """Thêm dòng Số: ... / Thành phố, ngày ..."""
    # Số - trái
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(p1, space_before=60, space_after=60, line_spacing=240)
    add_formatted_run(p1, "Số: ………", size=FONT_SIZE_META, bold=False)

    # Ngày tháng - phải (thêm vào cùng paragraph như tab, hoặc paragraph riêng)
    # Dùng paragraph mới căn phải
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(p2, space_before=-120, space_after=60, line_spacing=240)
    add_formatted_run(p2, "Thành phố Hà Nội, ngày … tháng … năm 20…", size=FONT_SIZE_META, bold=False)


def add_title_line(doc, title="THÔNG TƯ"):
    """Thêm dòng tiêu đề (ví dụ: Căn cứ Luật..., Hoặc tiêu đề chính)"""
    # Tìm xem doc có tiêu đề lớn (ALL CAPS, không phải Căn c�ứ) không
    # Nếu có thì giữ nguyên, không thêm mới
    # Chỉ thêm title nếu chưa có
    # Ở đây: đơn giản thêm 1 paragraph trống rồi để người dùng tự điền
    pass


def add_two_column_meta(doc):
    """Thêm dòng Số + Thành phố trên 1 hàng (trái | phải)"""
    table = doc.add_table(rows=1, cols=2)
    table.allow_autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Trái: Số:
    cell_l = table.cell(0, 0)
    cell_l.width = Cm(8.5)
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(p_l, space_before=60, space_after=60, line_spacing=240)
    add_formatted_run(p_l, "Số: ……………………", size=FONT_SIZE_META, bold=False)

    # Phải: Thành phố, ngày
    cell_r = table.cell(0, 1)
    cell_r.width = Cm(8.5)
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(p_r, space_before=60, space_after=60, line_spacing=240)
    add_formatted_run(p_r, "Hà Nội, ngày … tháng … năm 20…", size=FONT_SIZE_META, bold=False)

    return table


# --- Backward-compat aliases ---
def add_gov_header(doc):
    """Alias for build_header_table"""
    return build_header_table(doc)


# ==== MAIN ====
def process_file(filepath, dry_run=False, backup=False):
    """Xử lý 1 file: thêm header chuẩn"""
    filename = os.path.basename(filepath)
    print(f"\n{'='*60}")
    print(f"Xử lý: {filename}")
    print(f"{'='*60}")

    if backup and not dry_run:
        backup_dir = os.path.join(os.path.dirname(filepath), "_backup")
        os.makedirs(backup_dir, exist_ok=True)
        backup_path = os.path.join(backup_dir, filename)
        shutil.copy2(filepath, backup_path)
        print(f"  📦 Backup: {backup_path}")

    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  ❌ Lỗi mở file: {e}")
        return False

    # Kiểm tra xem đã có header chưa
    has_chx = False
    total_paras = len(doc.paragraphs)
    for para in doc.paragraphs:
        if total_paras > 50 and para._p.getparent() is not None:
            break
        if "XÃ HỘI CHỦ NGHĨA VIỆT NAM" in para.text.upper():
            has_chx = True
            break
        if "CHXHCNVN" in para.text.upper() or "CỘNG HÒA" in para.text.upper():
            has_chx = True
            break

    if has_chx:
        print(f"  ⚠️  Đã có header CHXHCNVN, bỏ qua")
        return True

    if dry_run:
        print(f"  ✅ [DRY-RUN] Sẽ thêm header chuẩn vào đầu file")
        return True

    # Lưu lại nội dung cũ để xóa sau
    # Cách: đọc tất cả content, xóa, thêm header mới rồi thêm lại content
    old_paras = []
    for p in doc.paragraphs:
        old_paras.append(p.text)

    # Xóa tất cả paragraphs
    body = doc.element.body
    # Giữ lại section properties
    sectPr = None
    for child in list(body):
        if child.tag.endswith("sectPr"):
            sectPr = child
            body.remove(child)
            break

    # Xóa tất cả elements trong body
    for child in list(body):
        body.remove(child)

    # Thêm header table
    build_header_table(doc)

    # Thêm dòng Số / Thành phố
    add_two_column_meta(doc)

    # Thêm paragraph trống
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, space_before=0, space_after=0, line_spacing=0)

    # Thêm lại nội dung cũ (trừ các dòng header cũ nếu có)
    for text in old_paras:
        # Bỏ qua dòng trống đầu tiên và dòng header cũ
        if not text.strip():
            p = doc.add_paragraph()
            set_paragraph_format(p, space_before=0, space_after=0, line_spacing=0)
            continue

        # Bỏ qua các dòng header cũ nếu có
        upper = text.strip().upper()
        skip_patterns = [
            "BỘ GIÁO DỤC VÀ ĐÀO TẠO",
            "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
            "ĐỘC LẬP - TỰ DO - HẠNH PHÚC",
            "Độc lập - Tự do - Hạnh phúc",
        ]
        if any(pattern in upper for pattern in skip_patterns):
            continue

        # Bỏ qua dòng "Số:" và dòng "Thành phố" cũ nếu có
        if upper.startswith("SỐ:") or upper.startswith("SỐ ") or "SỐ:" in upper[:15]:
            continue
        if upper.startswith("THÀNH PHỐ HÀ NỘI") or (upper.startswith("THÀNH PHỐ") and "NGÀY" in upper):
            continue

        p = doc.add_paragraph()
        # Căn giữa cho tiêu độethôn toàn văn bản
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(p, space_before=0, space_after=0, line_spacing=0)
        add_formatted_run(p, text, size=FONT_SIZE_META, bold=False)

    # Save
    try:
        doc.save(filepath)
        print(f"  ✅ Đã lưu file")
        return True
    except Exception as e:
        print(f"  ❌ Lỗi lưu file: {e}")
        return False


def parse_args():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--templates-dir",
        type=Path,
        default=Path(__file__).resolve().parent / "templates",
        help="Directory containing DOCX templates (defaults to repository templates/)",
    )
    parser.add_argument("--dry-run", action="store_true", help="Preview without writing files")
    backup_group = parser.add_mutually_exclusive_group()
    backup_group.add_argument("--backup", dest="backup", action="store_true")
    backup_group.add_argument("--no-backup", dest="backup", action="store_false")
    parser.set_defaults(backup=True)
    return parser.parse_args()


def main():
    args = parse_args()
    dry_run = args.dry_run
    backup = args.backup

    if dry_run:
        print("=== DRY RUN MODE - Không ghi file ===")

    templates_dir = args.templates_dir.expanduser().resolve()
    if not templates_dir.exists():
        print(f"❌ Không tìm thấy thư mục templates tại: {templates_dir}")
        return

    template_files = sorted(templates_dir.glob("*.docx"))
    print(f"\nTìm thấy {len(template_files)} file templates:")
    for f in template_files:
        print(f"  - {f.name}")

    success = 0
    skip = 0
    fail = 0

    for f in template_files:
        try:
            result = process_file(str(f), dry_run=dry_run, backup=backup)
            if result:
                success += 1
            else:
                skip += 1
        except Exception as e:
            print(f"  ❌ Lỗi xử lý {f.name}: {e}")
            fail += 1

    print(f"\n{'='*60}")
    print(f"KẾT QUẢ:")
    print(f"  ✅ Thành công: {success}")
    print(f"  ⚠️  Bỏ qua:     {skip}")
    print(f"  ❌ Lỗi:        {fail}")
    print(f"{'='*60}")

    if not dry_run and success > 0:
        print("\n✨ Đã xử lý xong! Mở các file templates xem kết quả nhé.")


if __name__ == "__main__":
    main()
